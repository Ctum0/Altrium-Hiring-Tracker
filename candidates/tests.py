from io import BytesIO

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse

from accounts.models import Role
from candidates.models import Candidate, JobApplication
from jobs.models import InterviewRound, Job
from notifications.models import Notification

User = get_user_model()


@override_settings(
    STORAGES={
        'default': {'BACKEND': 'django.core.files.storage.FileSystemStorage'},
        'staticfiles': {'BACKEND': 'django.contrib.staticfiles.storage.StaticFilesStorage'},
    }
)
class CandidatesBaseTestCase(TestCase):
    def setUp(self):
        self.hr = User.objects.create_user(
            username='hr', password='pass12345', role=Role.HR,
            first_name='Hana',
        )
        self.interviewer = User.objects.create_user(
            username='iv', password='pass12345', role=Role.INTERVIEWER,
            first_name='Ivan',
        )
        self.management = User.objects.create_user(
            username='mgmt', password='pass12345', role=Role.MANAGEMENT
        )
        self.job = Job.objects.create(title='Backend', created_by=self.hr)
        self.round1 = InterviewRound.objects.create(job=self.job, name='Screen', order=1)
        self.round2 = InterviewRound.objects.create(job=self.job, name='Tech', order=2)
        self.candidate = Candidate.objects.create(
            first_name='Ada', last_name='Lovelace',
            email='ada@example.com', skills='Python, Django',
        )
        self.application = JobApplication.objects.create(
            candidate=self.candidate, job=self.job,
            status=JobApplication.Status.NEW,
        )

    def login(self, username):
        assert self.client.login(username=username, password='pass12345')


class CandidateVisibilityTests(CandidatesBaseTestCase):
    def _add_other_candidate(self):
        other = Candidate.objects.create(
            first_name='Grace', last_name='Hopper', email='grace@example.com'
        )
        JobApplication.objects.create(candidate=other, job=self.job)
        return other

    def test_interviewer_sees_only_assigned(self):
        self._add_other_candidate()
        self.application.assigned_to = self.interviewer
        self.application.save()

        self.login('iv')
        r = self.client.get(reverse('candidates:list'))
        self.assertContains(r, 'Lovelace')
        self.assertNotContains(r, 'Hopper')

    def test_management_sees_all(self):
        self._add_other_candidate()
        self.login('mgmt')
        r = self.client.get(reverse('candidates:list'))
        self.assertContains(r, 'Hopper')

    def test_hr_sees_all(self):
        self._add_other_candidate()
        self.login('hr')
        r = self.client.get(reverse('candidates:list'))
        self.assertContains(r, 'Hopper')

    def test_interviewer_cannot_score(self):
        self.login('iv')
        r = self.client.post(reverse('candidates:score', args=[self.candidate.pk]), {'score': '90'})
        self.assertEqual(r.status_code, 302)
        self.candidate.refresh_from_db()
        self.assertIsNone(self.candidate.score)


class ScoreTests(CandidatesBaseTestCase):
    def test_score_valid(self):
        self.login('hr')
        r = self.client.post(reverse('candidates:score', args=[self.candidate.pk]), {'score': '85'})
        self.assertEqual(r.status_code, 302)
        self.candidate.refresh_from_db()
        self.assertEqual(self.candidate.score, 85)

    def test_score_boundaries(self):
        self.login('hr')
        # Valid boundary values are accepted.
        for value, expected in [('0', 0), ('100', 100)]:
            self.client.post(
                reverse('candidates:score', args=[self.candidate.pk]),
                {'score': value},
            )
            self.candidate.refresh_from_db()
            self.assertEqual(self.candidate.score, expected)
        # Invalid values are rejected (score unchanged from last valid).
        for value in ['-1', '101']:
            self.client.post(
                reverse('candidates:score', args=[self.candidate.pk]),
                {'score': value},
            )
            self.candidate.refresh_from_db()
            self.assertEqual(self.candidate.score, 100)

    def test_score_non_numeric(self):
        self.login('hr')
        self.client.post(reverse('candidates:score', args=[self.candidate.pk]), {'score': 'abc'})
        self.candidate.refresh_from_db()
        self.assertIsNone(self.candidate.score)

    def test_score_clear(self):
        self.candidate.score = 90
        self.candidate.save()
        self.login('hr')
        self.client.post(reverse('candidates:score', args=[self.candidate.pk]), {'score': ''})
        self.candidate.refresh_from_db()
        self.assertIsNone(self.candidate.score)


class AssignmentTests(CandidatesBaseTestCase):
    def test_assign_creates_notification(self):
        self.login('hr')
        r = self.client.post(
            reverse('candidates:assign', args=[self.application.pk]),
            {'interviewer': self.interviewer.pk},
        )
        self.assertEqual(r.status_code, 302)
        self.application.refresh_from_db()
        self.assertEqual(self.application.assigned_to, self.interviewer)
        self.assertTrue(Notification.objects.filter(recipient=self.interviewer).exists())

    def test_assign_missing_interviewer(self):
        self.login('hr')
        r = self.client.post(reverse('candidates:assign', args=[self.application.pk]), {})
        self.assertEqual(r.status_code, 302)
        self.application.refresh_from_db()
        self.assertIsNone(self.application.assigned_to)

    def test_assign_same_interviewer_no_dup_notification(self):
        self.application.assigned_to = self.interviewer
        self.application.save()
        self.login('hr')
        self.client.post(
            reverse('candidates:assign', args=[self.application.pk]),
            {'interviewer': self.interviewer.pk},
        )
        count = Notification.objects.filter(recipient=self.interviewer).count()
        self.assertEqual(count, 0)

    def test_interviewer_cannot_assign(self):
        self.login('iv')
        self.client.post(
            reverse('candidates:assign', args=[self.application.pk]),
            {'interviewer': self.interviewer.pk},
        )
        self.application.refresh_from_db()
        self.assertIsNone(self.application.assigned_to)


class UploadTests(CandidatesBaseTestCase):
    def _docx_file(self, name='cv.docx'):
        from docx import Document
        doc = Document()
        doc.add_paragraph('Jane Smith')
        doc.add_paragraph('Email: jane@example.com')
        doc.add_paragraph('Skills: Python, Django')
        bio = BytesIO()
        doc.save(bio)
        return SimpleUploadedFile(
            name, bio.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

    def test_upload_creates_candidate_and_application(self):
        self.login('hr')
        r = self.client.post(reverse('candidates:upload'), {
            'job': self.job.pk,
            'files': [self._docx_file()],
        })
        self.assertEqual(r.status_code, 302)
        self.assertTrue(
            Candidate.objects.filter(resume_text__contains='Jane Smith').exists()
        )
        self.assertTrue(
            JobApplication.objects.filter(
                job=self.job, candidate__resume_text__contains='Jane Smith'
            ).exists()
        )

    def test_upload_requires_job(self):
        self.login('hr')
        r = self.client.post(reverse('candidates:upload'), {
            'files': [self._docx_file()],
        })
        self.assertEqual(r.status_code, 200)  # re-rendered with error

    def test_upload_requires_files(self):
        self.login('hr')
        r = self.client.post(reverse('candidates:upload'), {'job': self.job.pk})
        self.assertEqual(r.status_code, 200)

    def test_upload_duplicate_email_dedupes(self):
        Candidate.objects.create(email='jane@example.com')
        self.login('hr')
        self.client.post(reverse('candidates:upload'), {
            'job': self.job.pk,
            'files': [self._docx_file()],
        })
        self.assertEqual(
            Candidate.objects.filter(email='jane@example.com').count(), 1
        )

    def test_upload_binary_file_no_crash(self):
        self.login('hr')
        garbage = SimpleUploadedFile(
            'bad.pdf',
            b'\x89PNG\r\n\x1a\nnot really a pdf' * 100,
            content_type='application/pdf',
        )
        r = self.client.post(reverse('candidates:upload'), {
            'job': self.job.pk,
            'files': [garbage],
        })
        self.assertEqual(r.status_code, 302)  # falls back gracefully, no 500
        self.assertTrue(Candidate.objects.filter(source='upload').exists())

    def test_upload_rejects_unsupported_extension(self):
        self.login('hr')
        bad = SimpleUploadedFile(
            'virus.exe', b'MZ\x90\x00', content_type='application/octet-stream'
        )
        r = self.client.post(reverse('candidates:upload'), {
            'job': self.job.pk,
            'files': [bad],
        })
        self.assertEqual(r.status_code, 200)  # re-rendered with error, not processed
        self.assertFalse(Candidate.objects.filter(source='upload').exists())

    def test_upload_requires_hr(self):
        self.login('iv')
        r = self.client.get(reverse('candidates:upload'))
        self.assertEqual(r.status_code, 302)


class ImportTests(CandidatesBaseTestCase):
    def _mock_parse(self, **overrides):
        from unittest.mock import patch
        parsed = {
            'first_name': 'Jane', 'last_name': 'Smith',
            'email': 'jane@example.com', 'phone': '', 'skills': ['React'],
        }
        parsed.update(overrides)
        return patch('candidates.views.parse_cv', return_value=parsed)

    def test_import_with_email(self):
        self.login('hr')
        with self._mock_parse():
            r = self.client.post(reverse('candidates:import'), {
                'job': self.job.pk,
                'source': 'LinkedIn',
                'profile_text': 'Jane Smith\nEmail: jane@example.com\nSkills: React, Node',
            })
        self.assertEqual(r.status_code, 302)
        self.assertTrue(Candidate.objects.filter(email='jane@example.com').exists())

    def test_import_duplicate_links_to_job(self):
        Candidate.objects.create(email='jane@example.com', first_name='Jane')
        self.login('hr')
        with self._mock_parse():
            self.client.post(reverse('candidates:import'), {
                'job': self.job.pk,
                'source': 'LinkedIn',
                'profile_text': 'Jane Smith\nEmail: jane@example.com',
            })
        self.assertEqual(
            Candidate.objects.filter(email='jane@example.com').count(), 1
        )
        cand = Candidate.objects.get(email='jane@example.com')
        self.assertTrue(cand.applications.filter(job=self.job).exists())

    def test_import_requires_hr(self):
        self.login('mgmt')
        r = self.client.get(reverse('candidates:import'))
        self.assertEqual(r.status_code, 302)


class SearchFilterTests(CandidatesBaseTestCase):
    def setUp(self):
        super().setUp()
        Candidate.objects.create(
            first_name='Grace', last_name='Hopper',
            email='grace@example.com', skills='COBOL',
        )
        JobApplication.objects.create(
            candidate=Candidate.objects.get(email='grace@example.com'),
            job=self.job,
        )

    def test_search_by_name(self):
        self.login('hr')
        r = self.client.get(reverse('candidates:list'), {'q': 'Lovelace'})
        self.assertContains(r, 'Lovelace')
        self.assertNotContains(r, 'Hopper')

    def test_search_by_skill(self):
        self.login('hr')
        r = self.client.get(reverse('candidates:list'), {'q': 'COBOL'})
        self.assertContains(r, 'Hopper')
        self.assertNotContains(r, 'Lovelace')

    def test_filter_by_stage(self):
        self.login('hr')
        r = self.client.get(reverse('candidates:list'), {'stage': 'new'})
        self.assertContains(r, 'Lovelace')

    def test_filter_by_job(self):
        self.login('hr')
        r = self.client.get(reverse('candidates:list'), {'job': self.job.pk})
        self.assertContains(r, 'Lovelace')

    def test_pagination_50_per_page(self):
        self.login('hr')
        for i in range(55):
            Candidate.objects.create(
                email=f'bulk{i}@example.com', first_name=f'Bulk{i}'
            )
            JobApplication.objects.create(
                candidate=Candidate.objects.get(email=f'bulk{i}@example.com'),
                job=self.job,
            )
        # Newest first: Bulk54 is on page 1, Bulk0 is on page 2.
        r = self.client.get(reverse('candidates:list'))
        self.assertContains(r, 'Bulk54')
        self.assertNotContains(r, 'Bulk0')
        r2 = self.client.get(reverse('candidates:list'), {'page': 2})
        self.assertContains(r2, 'Bulk0')

    def test_page_out_of_range_clamped(self):
        self.login('hr')
        r = self.client.get(reverse('candidates:list'), {'page': '999'})
        self.assertEqual(r.status_code, 200)
        r = self.client.get(reverse('candidates:list'), {'page': 'abc'})
        self.assertEqual(r.status_code, 200)
        r = self.client.get(reverse('candidates:list'), {'page': '-3'})
        self.assertEqual(r.status_code, 200)


class CandidateRbacSecurityTests(CandidatesBaseTestCase):
    def test_unassigned_interviewer_blocked_from_candidate_detail(self):
        # Interviewer is NOT assigned to application
        self.login('iv')
        url = reverse('candidates:detail', kwargs={'pk': self.candidate.pk})
        response = self.client.get(url)
        self.assertEqual(response.status_code, 403)

    def test_assigned_interviewer_can_view_candidate_detail(self):
        # Assign candidate application to interviewer
        self.application.assigned_to = self.interviewer
        self.application.save()
        self.login('iv')
        url = reverse('candidates:detail', kwargs={'pk': self.candidate.pk})
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)

    def test_unassigned_interviewer_blocked_from_ai_fit_assessment(self):
        self.login('iv')
        url = reverse('candidates:ai_fit', kwargs={'pk': self.application.pk})
        response = self.client.post(url)
        self.assertEqual(response.status_code, 403)

    def test_interviewer_blocked_from_hr_score_update(self):
        self.login('iv')
        url = reverse('candidates:score', kwargs={'pk': self.candidate.pk})
        response = self.client.post(url, {'score': '95'}, follow=True)
        # Should redirect or error, score remains unchanged
        self.candidate.refresh_from_db()
        self.assertNotEqual(self.candidate.score, 95)

    def test_management_blocked_from_modifications(self):
        self.login('mgmt')
        url = reverse('candidates:score', kwargs={'pk': self.candidate.pk})
        response = self.client.post(url, {'score': '95'})
        self.assertEqual(response.status_code, 403)
        self.candidate.refresh_from_db()
        self.assertNotEqual(self.candidate.score, 95)
