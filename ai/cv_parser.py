"""Extract raw text from uploaded CV files (PDF / DOCX)."""
import io
import logging

logger = logging.getLogger(__name__)


def extract_text(uploaded_file) -> str:
    """Return raw text from a PDF or DOCX file object."""
    name = (uploaded_file.name or '').lower()
    data = uploaded_file.read()
    if name.endswith('.pdf'):
        return _from_pdf(data)
    if name.endswith('.docx'):
        return _from_docx(data)
    # Fallback: try to read as text
    try:
        return data.decode('utf-8', errors='replace')
    except Exception:
        return ''


def _from_pdf(data: bytes) -> str:
    try:
        from pdfminer.high_level import extract_text as pdf_extract
        return pdf_extract(io.BytesIO(data))
    except Exception as exc:
        logger.error('PDF extraction failed: %s', exc)
        return ''


def _from_docx(data: bytes) -> str:
    try:
        import docx
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(' | '.join(cells))
        return '\n'.join(parts)
    except Exception as exc:
        logger.error('DOCX extraction failed: %s', exc)
        return ''
